from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SCRIPT_DIR = Path(__file__).resolve().parent
RESULTS_DIR = SCRIPT_DIR / "results"
FIGURES_DIR = SCRIPT_DIR / "figures"
PREP_SUMMARY_PATH = RESULTS_DIR / "prep_summary.json"
BASELINE_METRICS_PATH = RESULTS_DIR / "baseline_metrics.json"
BASELINE_COMPARISON_PATH = RESULTS_DIR / "baseline_comparison.csv"
OUTPUT_DOC = SCRIPT_DIR / "Milestone4_deliverable.docx"


def set_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def style_paragraph(paragraph, align=None, after=6, line=1.15):
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_para(
    doc,
    text,
    *,
    size=12,
    bold=False,
    italic=False,
    align=None,
    after=6,
    line=1.15,
    style=None,
):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    style_paragraph(paragraph, align=align, after=after, line=line)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    set_font(run, size=14 if level == 1 else 12, bold=True)
    style_paragraph(paragraph, after=6, line=1.1)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_font(run)
    style_paragraph(paragraph, after=3, line=1.1)
    return paragraph


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = str(header)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_font(run, size=11)
                style_paragraph(paragraph, after=0, line=1.0)
    return table


def add_figure(doc, path: Path, caption: str, width=6.2):
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_paragraph.add_run(caption)
    set_font(run, size=11, italic=True)
    style_paragraph(caption_paragraph, after=8, line=1.0)


def metric(value: float) -> str:
    return f"{value:.3f}"


def pct(value: float) -> str:
    return f"{value:.2f}%"


def build_task_table(comparison: pd.DataFrame, task_name: str, problem_type: str) -> list[list[str]]:
    subset = comparison.loc[comparison["task"] == task_name].copy()
    subset = subset.sort_values("selected_by_validation", ascending=False)
    rows = []
    for _, row in subset.iterrows():
        selected = "Yes" if row["selected_by_validation"] else "No"
        if problem_type == "multiclass":
            rows.append(
                [
                    row["model"],
                    selected,
                    metric(row["validation_macro_f1"]),
                    metric(row["test_accuracy"]),
                    metric(row["test_macro_f1"]),
                    metric(row["test_balanced_accuracy"]),
                ]
            )
        else:
            rows.append(
                [
                    row["model"],
                    selected,
                    metric(row["validation_f1"]),
                    metric(row["test_accuracy"]),
                    metric(row["test_f1"]),
                    metric(row["test_roc_auc"]),
                ]
            )
    return rows


def build():
    prep_summary = json.loads(PREP_SUMMARY_PATH.read_text())
    baseline_metrics = json.loads(BASELINE_METRICS_PATH.read_text())
    comparison = pd.read_csv(BASELINE_COMPARISON_PATH)

    uci_multi = baseline_metrics["tasks"]["uci_multiclass"]
    uci_binary = baseline_metrics["tasks"]["uci_binary_early"]
    oulad_binary = baseline_metrics["tasks"]["oulad_binary_early"]

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    for style_name in ["Heading 1", "Heading 2", "Heading 3", "List Bullet"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_para(doc, "FACULTY OF BUSINESS AND IT", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, "MBAI 5600G - Applied Integrative Analytics Capstone Project", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, "Milestone 4 Report", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    add_para(
        doc,
        "AI-Based Student Dropout Early Warning System for Higher Education:\nReproduction and Cross-Dataset Generalization on UCI and OULAD",
        size=14,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=12,
    )
    add_para(doc, "Student Group Name: [Insert Group Name]", align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "Team Members: Xinyu Wang, Wu Hanming", align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "Submission: Milestone 4 Data Preprocessing and Baseline Modeling", align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "Due Date: June 23, 2026", align=WD_ALIGN_PARAGRAPH.CENTER, after=4)

    doc.add_section(WD_SECTION.NEW_PAGE)

    add_heading(doc, "1. Updated Executive Summary", level=1)
    exec_summary = (
        "This milestone converted the project from exploratory analysis into a reproducible modeling workflow across the "
        "UCI and OULAD datasets. The preprocessing stage produced three modeling tables: a UCI multiclass reproduction set, "
        "a leakage-aware UCI binary early-warning set, and a student-level OULAD early-warning set built from multi-table "
        "assessment and Virtual Learning Environment activity. Data cleaning confirmed zero duplicates in all final tables, "
        "zero missing values in UCI, and limited retained missingness in OULAD, primarily imd_band at 3.41 percent. Missing "
        "values were handled inside the modeling pipeline by median imputation for numeric fields and most-frequent imputation "
        "for categorical fields. Numeric outliers were detected by an IQR review on the training split and treated through "
        "1st and 99th percentile clipping so extreme but plausible student behavior could be retained. The baseline evaluation "
        "used stratified 60-20-20 train-validation-test splitting, plus 5-fold cross-validation on the selected model for each "
        "task. For the direct UCI reproduction benchmark, XGBoost was the strongest article-aligned baseline with test accuracy "
        f"of {metric(uci_multi['best_test_metrics']['accuracy'])} and macro F1 of {metric(uci_multi['best_test_metrics']['macro_f1'])}. "
        f"For practical early warning, logistic regression remained strongest on UCI with F1 of {metric(uci_binary['best_test_metrics']['f1'])}, "
        f"while XGBoost was strongest on OULAD with F1 of {metric(oulad_binary['best_test_metrics']['f1'])} and ROC AUC of "
        f"{metric(oulad_binary['best_test_metrics']['roc_auc'])}. These results establish a documented and reproducible baseline "
        "for later optimization, transfer analysis, and explanation-stability work in Milestone 5."
    )
    add_para(doc, exec_summary)

    add_heading(doc, "2. Data Preprocessing Report", level=1)
    add_heading(doc, "2.1 Preprocessing Workflow Overview", level=2)
    add_para(
        doc,
        "The Milestone 4 preprocessing workflow was designed to transform the raw datasets collected in Milestone 3 into "
        "modeling-ready, leakage-aware analysis tables. UCI required only light restructuring because it was already a single "
        "clean CSV file. OULAD required a more extensive student-level build process because the relevant information was spread "
        "across studentInfo, studentRegistration, courses, assessments, studentAssessment, vle, and studentVle. The final output "
        "of preprocessing was three datasets: one direct UCI multiclass benchmark for reproduction, one UCI early-warning binary "
        "dataset with second-semester leakage removed, and one OULAD early-warning binary dataset aggregated to the student-module-presentation level."
    )
    add_bullet(doc, "UCI multiclass reproduction retained the original three-class Target variable: Dropout, Enrolled, and Graduate.")
    add_bullet(doc, "UCI binary early warning created is_attrition and removed all second-semester variables to avoid using late-course information.")
    add_bullet(doc, "OULAD binary early warning defined Withdrawn as the positive attrition class and aggregated only activity observed within the first 75 days.")

    add_heading(doc, "2.2 Data Cleaning Procedures", level=2)
    add_para(
        doc,
        "Data cleaning focused on duplicate checks, schema consistency, dataset consolidation, and invalid-information removal. "
        "No duplicate rows were found in the final UCI or OULAD modeling tables. OULAD consolidation used code_module, "
        "code_presentation, and id_student as the primary join keys. date_unregistration was explicitly removed from the modeling "
        "inputs because it directly reveals withdrawal timing and would create severe target leakage. In the OULAD build, all "
        "aggregate early-course features that were structurally absent because a student had no qualifying event were filled with zero, "
        "which reflects no early recorded activity rather than missing information."
    )
    add_table(
        doc,
        ["Dataset", "Rows", "Columns", "Duplicate Rows", "Missing Cells", "Modeling Role"],
        [
            [
                "UCI multiclass",
                prep_summary["uci"]["shape"][0],
                prep_summary["uci"]["shape"][1],
                prep_summary["uci"]["duplicate_total"],
                prep_summary["uci"]["missing_total"],
                "Direct reproduction benchmark",
            ],
            [
                "UCI binary early warning",
                prep_summary["uci"]["shape"][0],
                32,
                prep_summary["uci"]["duplicate_total"],
                prep_summary["uci"]["missing_total"],
                "Intervention-oriented extension",
            ],
            [
                "OULAD binary early warning",
                prep_summary["oulad"]["shape"][0],
                prep_summary["oulad"]["shape"][1],
                prep_summary["oulad"]["duplicate_total"],
                prep_summary["oulad"]["missing_total"],
                "Cross-dataset extension baseline",
            ],
        ],
    )

    add_heading(doc, "2.3 Missing Value Handling and Data Quality Improvements", level=2)
    add_para(
        doc,
        "The UCI datasets contained no missing values, so no row or variable removal was required during preprocessing. "
        "In contrast, the OULAD modeling table retained limited missingness in a small number of fields after consolidation. "
        "Those fields were retained because their missingness was modest and because removing the affected students would have "
        "reduced sample coverage unnecessarily. Categorical missingness was handled inside the modeling pipeline with most-frequent "
        "imputation, while numeric missingness was handled with median imputation. This approach preserves the full cohort while "
        "keeping preprocessing reproducible inside the train-only pipeline fit."
    )
    oulad_missing_rows = [
        [column, pct(value)]
        for column, value in prep_summary["oulad"]["missing_by_column_pct"].items()
    ]
    add_table(
        doc,
        ["OULAD Variable", "Missing Percentage"],
        oulad_missing_rows,
    )

    add_heading(doc, "2.4 Outlier Detection and Treatment", level=2)
    add_para(
        doc,
        "Outlier review was performed on the numeric training data using the IQR rule. The purpose was not to discard unusual students, "
        "but to identify whether extreme numeric values could distort linear and boosting baselines. Because unusually high or low "
        "engagement may still be educationally meaningful, the project did not delete outlier rows. Instead, numeric variables were "
        "winsorized inside the pipeline by clipping them to the 1st and 99th percentile range estimated from the training data only. "
        "This prevents leakage while reducing undue leverage from very large values."
    )
    outlier_rows = [
        [
            item["feature"],
            item["count"],
            pct(item["pct_rows"]),
        ]
        for item in uci_multi["outlier_summary"]["top_features"][:5]
    ]
    add_table(
        doc,
        ["Illustrative UCI Numeric Feature", "Flagged Cells", "Percent of Training Rows"],
        outlier_rows,
    )

    add_heading(doc, "3. Feature Engineering and Data Preparation", level=1)
    add_heading(doc, "3.1 Engineered Features and Feature Families", level=2)
    add_para(
        doc,
        "Feature preparation differed by dataset. In UCI, the main preprocessing choice was to preserve code-based categorical fields "
        "as categorical variables during modeling instead of leaving them as ordinary continuous numbers. In OULAD, the larger task "
        "was to derive student-level early-course summaries from event and assessment tables. Assessment features included early "
        "submission counts, mean and standard deviation of scores, weighted-score totals, submission delay, late-submission counts, "
        "and counts by assessment type. Virtual Learning Environment features included total clicks, event counts, active days, unique "
        "sites visited, and activity-type click totals. Registration timing, course length, previous attempts, education level, and "
        "other student descriptors were also retained."
    )
    add_bullet(doc, "Categorical encoding: one-hot encoding with unknown-category protection.")
    add_bullet(doc, "Numeric scaling: StandardScaler for logistic regression only; tree and boosting baselines used imputed and clipped raw numeric values.")
    add_bullet(doc, "Dimensionality reduction: not applied in Milestone 4 because the goal was reproducible baseline comparison rather than compressed representation learning.")
    add_bullet(doc, "Feature selection: no manual feature removal beyond leakage control, because the baseline stage should preserve the information environment seen by the chosen article models.")

    add_heading(doc, "3.2 Shared Feature Schema for Cross-Dataset Generalization", level=2)
    add_para(
        doc,
        "A shared feature schema was exported as both JSON and CSV so UCI and OULAD can later be compared at the concept level rather than the raw-column level. "
        "The schema groups variables into common families such as demographics, prior preparation, program setup, financial support, early academic progress, "
        "early engagement, and macro context. This is important because the two datasets do not share identical columns. UCI emphasizes structured academic "
        "and administrative variables, while OULAD adds rich temporal behavior through the Virtual Learning Environment. The shared schema therefore serves as "
        "the bridge for later cross-dataset transfer analysis and explanation-stability evaluation."
    )
    add_table(
        doc,
        ["Feature Family", "UCI Example", "OULAD Example"],
        [
            ["Demographics", "Gender, Age at enrollment", "gender, age_band, imd_band"],
            ["Prior preparation", "Previous qualification, parent qualifications", "highest_education, num_of_prev_attempts"],
            ["Program setup", "Course, Application mode", "code_module, code_presentation, studied_credits"],
            ["Early academic progress", "1st semester approvals and grades", "assessment_submission_ratio_early, assessment_score_mean_early"],
            ["Early engagement", "Not available in raw UCI", "vle_total_clicks_early, vle_active_days_early"],
        ],
    )

    add_heading(doc, "4. Data Partitioning Strategy", level=1)
    add_para(
        doc,
        "All three tasks used the same reproducible partitioning logic. First, each dataset was split into a stratified train-validation block "
        "and a hold-out test block. Second, the train-validation block was split again into separate train and validation sets. The resulting ratios "
        "were 60 percent train, 20 percent validation, and 20 percent test. Validation metrics were used to choose the strongest baseline model for "
        "each task. After model selection, the selected model was additionally checked with 5-fold stratified cross-validation on the combined train-validation "
        "block to assess stability. The test set remained untouched until the final baseline evaluation. This design reduces the risk of leakage and makes the "
        "reported test metrics more defensible than a single train-test split."
    )
    add_table(
        doc,
        ["Task", "Train Rows", "Validation Rows", "Test Rows", "Selection Metric", "CV Folds"],
        [
            [
                uci_multi["display_name"],
                uci_multi["partition_strategy"]["train_rows"],
                uci_multi["partition_strategy"]["validation_rows"],
                uci_multi["partition_strategy"]["test_rows"],
                uci_multi["model_selection_metric"],
                uci_multi["selected_model_cross_validation"]["folds"],
            ],
            [
                uci_binary["display_name"],
                uci_binary["partition_strategy"]["train_rows"],
                uci_binary["partition_strategy"]["validation_rows"],
                uci_binary["partition_strategy"]["test_rows"],
                uci_binary["model_selection_metric"],
                uci_binary["selected_model_cross_validation"]["folds"],
            ],
            [
                oulad_binary["display_name"],
                oulad_binary["partition_strategy"]["train_rows"],
                oulad_binary["partition_strategy"]["validation_rows"],
                oulad_binary["partition_strategy"]["test_rows"],
                oulad_binary["model_selection_metric"],
                oulad_binary["selected_model_cross_validation"]["folds"],
            ],
        ],
    )

    add_heading(doc, "5. Baseline Model Development and Evaluation", level=1)
    add_heading(doc, "5.1 Baseline Reproduction Design", level=2)
    add_para(
        doc,
        "The baseline reproduction benchmark was based on the workflow described by Islam et al. (2025) [1]. The article reports decision tree, random forest, "
        "gradient boosting, and XGBoost as the main machine-learning baselines, with XGBoost reported as the strongest model at 83 percent accuracy. To align the "
        "Milestone 4 reproduction with that paper, all four article-aligned models were implemented for the UCI benchmark, and logistic regression was retained as "
        "an additional reference baseline for early-warning comparisons. Because the original paper focuses on the UCI benchmark problem, the UCI multiclass task is "
        "the direct reproduction setting. The UCI binary early-warning task and the OULAD binary early-warning task are project extensions rather than one-to-one reproductions."
    )
    add_table(
        doc,
        ["Model", "Role in This Milestone", "Key Settings"],
        [
            ["logistic_regression", "reference baseline", "balanced class weights, max_iter=3000, numeric scaling"],
            ["decision_tree", "article baseline", "max_depth=8, min_samples_leaf=10, balanced class weights"],
            ["random_forest", "article baseline", "400 trees, min_samples_leaf=2, balanced subsampling"],
            ["gradient_boosting", "article baseline", "200 estimators, learning_rate=0.05, depth=3, subsample=0.9"],
            ["xgboost", "article baseline", "250 estimators, depth=5, learning_rate=0.05, tree_method=hist"],
        ],
    )

    add_heading(doc, "5.2 UCI Reproduction Results", level=2)
    add_para(
        doc,
        "The UCI multiclass benchmark selected XGBoost as the strongest article-aligned model on validation macro F1, which also matches the article's reported best model family. "
        f"On the hold-out test set, the reproduced XGBoost baseline achieved accuracy {metric(uci_multi['best_test_metrics']['accuracy'])}, macro F1 "
        f"{metric(uci_multi['best_test_metrics']['macro_f1'])}, and balanced accuracy {metric(uci_multi['best_test_metrics']['balanced_accuracy'])}. "
        f"This is {abs(uci_multi['article_alignment']['accuracy_gap_vs_reported']) * 100:.2f} percentage points below the 83 percent accuracy reported in [1]. "
        "The discrepancy is reasonable at this stage because the current reproduction prioritizes a transparent, leakage-aware baseline pipeline rather than full article-level optimization. "
        "Differences in preprocessing detail, hyperparameter tuning, or experimental setup can materially affect multiclass performance on this dataset."
    )
    add_table(
        doc,
        ["Model", "Selected", "Validation Macro F1", "Test Accuracy", "Test Macro F1", "Test Balanced Acc."],
        build_task_table(comparison, "uci_multiclass", "multiclass"),
    )
    add_figure(
        doc,
        FIGURES_DIR / "uci_multiclass_xgboost_confusion_matrix.png",
        "Figure 1. UCI multiclass XGBoost confusion matrix for the selected reproduction baseline.",
    )
    add_figure(
        doc,
        FIGURES_DIR / "uci_multiclass_xgboost_top_features.png",
        "Figure 2. UCI multiclass XGBoost top features show that semester academic progress remains the dominant signal family.",
    )

    add_heading(doc, "5.3 Early-Warning Extension Results", level=2)
    add_para(
        doc,
        "The early-warning extensions answer a more practical intervention question than the direct reproduction benchmark. In the UCI binary early-warning task, "
        "logistic regression remained the strongest validation model. This result is notable because it was achieved after removing all second-semester variables, which makes the task more realistic for real intervention timing. "
        f"The final UCI binary test results were accuracy {metric(uci_binary['best_test_metrics']['accuracy'])}, F1 {metric(uci_binary['best_test_metrics']['f1'])}, and ROC AUC {metric(uci_binary['best_test_metrics']['roc_auc'])}. "
        "In OULAD, the selected baseline shifted to XGBoost, which suggests that the richer behavioral and temporal feature space benefits from a stronger nonlinear learner. "
        f"The final OULAD test results were accuracy {metric(oulad_binary['best_test_metrics']['accuracy'])}, F1 {metric(oulad_binary['best_test_metrics']['f1'])}, and ROC AUC {metric(oulad_binary['best_test_metrics']['roc_auc'])}."
    )
    add_table(
        doc,
        ["Model", "Selected", "Validation F1", "Test Accuracy", "Test F1", "Test ROC AUC"],
        build_task_table(comparison, "uci_binary_early", "binary"),
    )
    add_table(
        doc,
        ["Model", "Selected", "Validation F1", "Test Accuracy", "Test F1", "Test ROC AUC"],
        build_task_table(comparison, "oulad_binary_early", "binary"),
    )
    add_figure(
        doc,
        FIGURES_DIR / "uci_binary_early_logistic_regression_top_features.png",
        "Figure 3. UCI binary early-warning logistic regression top features from the selected intervention-oriented baseline.",
    )
    add_figure(
        doc,
        FIGURES_DIR / "oulad_binary_early_xgboost_confusion_matrix.png",
        "Figure 4. OULAD early-warning XGBoost confusion matrix for the selected cross-dataset extension baseline.",
    )
    add_figure(
        doc,
        FIGURES_DIR / "oulad_binary_early_xgboost_top_features.png",
        "Figure 5. OULAD early-warning XGBoost top features emphasize assessment behavior, course context, and early engagement.",
    )

    add_heading(doc, "5.4 Cross-Validation Stability of Selected Models", level=2)
    add_para(
        doc,
        "After model selection, the chosen baseline for each task was rechecked with 5-fold stratified cross-validation on the combined training and validation data. "
        "The CV results suggest that the selected models are reasonably stable rather than relying on one lucky split. For example, the selected UCI multiclass XGBoost baseline "
        f"produced mean CV accuracy {metric(uci_multi['selected_model_cross_validation']['metrics']['accuracy']['mean'])} and mean CV macro F1 "
        f"{metric(uci_multi['selected_model_cross_validation']['metrics']['macro_f1']['mean'])}. The selected UCI binary logistic baseline produced mean CV F1 "
        f"{metric(uci_binary['selected_model_cross_validation']['metrics']['f1']['mean'])}, while the selected OULAD XGBoost baseline produced mean CV F1 "
        f"{metric(oulad_binary['selected_model_cross_validation']['metrics']['f1']['mean'])}. This supports using these models as defensible references for Milestone 5."
    )
    add_table(
        doc,
        ["Selected Task Baseline", "Primary Metric", "CV Mean", "CV Std", "Supporting Metric", "Supporting CV Mean"],
        [
            [
                f"{uci_multi['display_name']} ({uci_multi['best_model']})",
                "macro_f1",
                metric(uci_multi["selected_model_cross_validation"]["metrics"]["macro_f1"]["mean"]),
                metric(uci_multi["selected_model_cross_validation"]["metrics"]["macro_f1"]["std"]),
                "accuracy",
                metric(uci_multi["selected_model_cross_validation"]["metrics"]["accuracy"]["mean"]),
            ],
            [
                f"{uci_binary['display_name']} ({uci_binary['best_model']})",
                "f1",
                metric(uci_binary["selected_model_cross_validation"]["metrics"]["f1"]["mean"]),
                metric(uci_binary["selected_model_cross_validation"]["metrics"]["f1"]["std"]),
                "roc_auc",
                metric(uci_binary["selected_model_cross_validation"]["metrics"]["roc_auc"]["mean"]),
            ],
            [
                f"{oulad_binary['display_name']} ({oulad_binary['best_model']})",
                "f1",
                metric(oulad_binary["selected_model_cross_validation"]["metrics"]["f1"]["mean"]),
                metric(oulad_binary["selected_model_cross_validation"]["metrics"]["f1"]["std"]),
                "roc_auc",
                metric(oulad_binary["selected_model_cross_validation"]["metrics"]["roc_auc"]["mean"]),
            ],
        ],
    )

    add_heading(doc, "6. Baseline Modeling Assessment and Future Optimization Plan", level=1)
    add_para(
        doc,
        "Milestone 4 established a reproducible baseline, but it also exposed the main improvement paths for Milestone 5. First, the UCI reproduction result remains below the "
        "accuracy reported by Islam et al. (2025) [1], so the next step is to test whether targeted hyperparameter tuning or more article-specific preprocessing choices narrow the gap. "
        "Second, the UCI binary and OULAD binary tasks now provide a clearer intervention-oriented benchmark, but they still need stronger robustness checks such as alternative cutoff windows, "
        "course-aware validation, and sensitivity tests around feature families. Third, the selected OULAD baseline still uses course identifiers as part of the feature space, which is acceptable for a baseline "
        "but should be stress-tested when generalization becomes the main goal. Finally, the current explainability artifacts are still feature-importance based rather than a full SHAP-oriented comparison, "
        "so the next milestone should extend beyond prediction quality into explanation quality and cross-dataset explanation stability."
    )
    add_bullet(doc, "Planned Milestone 5 activity 1: tune the selected baseline models, especially UCI multiclass XGBoost and OULAD early-warning XGBoost.")
    add_bullet(doc, "Planned Milestone 5 activity 2: test stronger validation schemes, including robustness across module-presentations and alternative early-course windows.")
    add_bullet(doc, "Planned Milestone 5 activity 3: begin explicit cross-dataset generalization and transferability analysis using the shared feature schema.")
    add_bullet(doc, "Planned Milestone 5 activity 4: add SHAP-based explanation analysis so performance and interpretability can be compared together.")

    add_heading(doc, "7. References", level=1)
    references = [
        '[1] M. M. Islam, F. H. Sojib, M. F. H. Mihad, M. Hasan, and M. Rahman, "The integration of explainable AI in Educational Data Mining for student academic performance prediction and support system," Telematics and Informatics Reports, vol. 18, Art. no. 100203, 2025, doi: 10.1016/j.teler.2025.100203.',
        '[2] V. Realinho, J. Machado, L. Baptista, and M. V. Martins, "Predicting Student Dropout and Academic Success," Data, vol. 7, no. 11, Art. no. 146, 2022, doi: 10.3390/data7110146.',
        '[3] V. Realinho, M. V. Martins, J. Machado, and L. Baptista, Predict Students\' Dropout and Academic Success [Dataset], UCI Machine Learning Repository, 2021. doi: 10.24432/C5MC89.',
        '[4] J. Kuzilek, M. Hlosta, and Z. Zdrahal, "Open University Learning Analytics dataset," Scientific Data, vol. 4, Art. no. 170171, 2017, doi: 10.1038/sdata.2017.171.',
        '[5] C. R. Cirak, H. Akilli, and Y. Ekinci, "Development of an early warning system for higher education institutions by predicting first-year student academic performance," Higher Education Quarterly, vol. 78, no. 4, 2024, doi: 10.1111/hequ.12539.',
        '[6] M. Phan, A. De Caigny, and K. Coussement, "A decision support framework to incorporate textual data for early student dropout prediction in higher education," Decision Support Systems, vol. 168, Art. no. 113940, 2023, doi: 10.1016/j.dss.2023.113940.',
    ]
    for ref in references:
        add_para(doc, ref, after=3, line=1.05)

    doc.save(OUTPUT_DOC)
    print(f"Wrote {OUTPUT_DOC}")


if __name__ == "__main__":
    build()
