from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SCRIPT_DIR = Path(__file__).resolve().parent
MILESTONE_DIR = SCRIPT_DIR.parent
SUMMARY_PATH = MILESTONE_DIR / "data" / "analysis" / "milestone3_eda_summary.json"
WORK_DOC = MILESTONE_DIR / "Milestone3_deliverable.docx"
ALT_WORK_DOC = MILESTONE_DIR / "Milestone3_deliverable_prof_style.docx"


def set_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def style_paragraph(p, align=None, after=6, line=1.15):
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align


def add_para(doc, text, *, size=12, bold=False, italic=False, align=None, after=6, line=1.15, style=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    style_paragraph(p, align=align, after=after, line=line)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    set_font(run, size=14 if level == 1 else 12, bold=True)
    style_paragraph(p, after=6, line=1.1)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_font(run)
    style_paragraph(p, after=3, line=1.1)
    return p


def pct(count, total):
    return round((count / total) * 100, 2) if total else 0


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_font(run)
                style_paragraph(p, after=0, line=1.0)
    return table


def add_figure(doc, path, caption, width=6.2):
    doc.add_picture(str(path), width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_font(run, size=11, italic=True)
    style_paragraph(cap, after=8, line=1.0)


def build():
    summary = json.loads(SUMMARY_PATH.read_text())
    uci = summary["uci"]
    oulad = summary["oulad"]

    uci_total = uci["shape"][0]
    uci_target = uci["target_counts"]
    oulad_total = oulad["table_shapes"]["studentInfo"][0]
    oulad_target = oulad["target_counts"]

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    for style_name in ["Heading 1", "Heading 2", "Heading 3", "List Bullet"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # Title page
    add_para(doc, "FACULTY OF BUSINESS AND IT", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, "MBAI 5600G - Applied Integrative Analytics Capstone Project", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, "Milestone 3 Report", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    add_para(
        doc,
        "AI-Based Student Dropout Early Warning System for Higher Education:\nReproduction and Cross-Dataset Generalization on UCI and OULAD",
        size=14,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=12,
    )
    add_para(doc, "Student Group Name: [Insert Group Name]", align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "Team Members: Xinyu Wang, Wu Hanming", align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "Submission: Milestone 3 Data Collection and Initial EDA", align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "Due Date: June 9, 2026", align=WD_ALIGN_PARAGRAPH.CENTER, after=4)

    doc.add_section(WD_SECTION.NEW_PAGE)

    add_heading(doc, "1. Updated Executive Summary", level=1)
    exec_summary = (
        "This milestone focused on collecting the two project datasets and performing initial exploratory data analysis "
        "before baseline modeling. First, the UCI Predict Students' Dropout and Academic Success dataset was finalized "
        "as the reproduction and benchmark dataset. It contains 4,424 records and 37 variables, with no missing values "
        "and no duplicate rows. Next, OULAD was collected as the primary evaluation dataset and organized as seven linked "
        "CSV files, including 32,593 student records, 173,912 assessment records, and more than 10.6 million raw Virtual Learning Environment (VLE) "
        "interaction rows. The initial EDA showed a clear pattern in both datasets. In UCI, dropout is most strongly "
        "associated with lower approved curricular units and lower semester grades. In OULAD, withdrawn students show much "
        "lower assessment participation, lower average scores, and much lower Virtual Learning Environment activity than pass or distinction groups. "
        "The data quality review also highlighted the main preparation issues for the next milestone, especially structural "
        "missingness, multi-table integration, and leakage control in OULAD. Overall, the milestone confirms that the project "
        "is ready to move into preprocessing, feature engineering, and baseline modeling in Milestone 4."
    )
    add_para(doc, exec_summary)

    add_heading(doc, "2. Data Collection and Dataset Description", level=1)
    add_heading(doc, "2.1 Finalized Dataset Sources and Acquisition Procedures", level=2)
    add_para(
        doc,
        "Two public datasets were collected for this capstone. The first is the UCI Predict Students' Dropout and "
        "Academic Success dataset, which is being used as the reproduction and benchmark dataset for the base paper "
        "by Islam et al. (2025) [1]-[3]. The second is OULAD, which is being used as the primary evaluation dataset "
        "because it provides richer behavioural and temporal information for studying model transferability and robustness [4]. "
        "The UCI dataset was downloaded as a single CSV file, while OULAD was downloaded as the official ZIP archive and "
        "extracted into seven CSV files. The datasets were stored locally in a structured project folder so that later "
        "preprocessing and modeling steps can be reproduced consistently."
    )
    add_para(
        doc,
        "No external API collection was required for this milestone. The first verification step was to confirm file access, "
        "row counts, major schema elements, target variables, and the presence of missing values or duplicates in the key tables. "
        "UCI required no integration during collection. OULAD required inspection of its linked-table design and preparation for "
        "later joins using code_module, code_presentation, and id_student."
    )

    add_heading(doc, "2.2 Dataset Structure and Characteristics", level=2)
    add_table(
        doc,
        ["Dataset", "Role", "Format", "Observations / Tables", "Target Variable", "Primary Data Type"],
        [
            [
                "UCI Predict Students' Dropout and Academic Success",
                "Reproduction and benchmark dataset",
                "Single CSV",
                "4,424 records / 37 columns",
                "Target: Graduate, Dropout, Enrolled",
                "Structured tabular data",
            ],
            [
                "OULAD",
                "Primary evaluation dataset",
                "ZIP with 7 CSV files",
                "32,593 student records; 173,912 assessment rows; 10,655,280 raw Virtual Learning Environment event rows",
                "final_result: Pass, Withdrawn, Fail, Distinction",
                "Structured multi-table learning analytics data",
            ],
        ],
    )
    add_para(
        doc,
        "The UCI dataset is compact and already easy to load for modeling, although many of its numeric fields are coded "
        "categorical or ordinal variables. OULAD is more complex and better reflects a realistic learning analytics environment. "
        "It includes demographic data, registration history, assessment records, and event-level Virtual Learning Environment interaction data across the "
        "2013B, 2013J, 2014B, and 2014J module-presentations."
    )

    add_heading(doc, "2.3 Variable Description and Data Integration Logic", level=2)
    add_para(
        doc,
        "For UCI, the most relevant variable families are demographic and background variables, academic preparation "
        "variables, institutional and financial indicators, and first- and second-semester academic performance variables. "
        "The strongest candidate predictors for later modeling are semester-approved curricular units, semester grades, "
        "admission grade, prior qualification grade, age at enrollment, debtor status, tuition payment status, and scholarship status."
    )
    add_para(
        doc,
        "For OULAD, the important variables are spread across multiple tables. The studentInfo table contains the core student "
        "and target information. studentRegistration contains enrollment and withdrawal timing fields. studentAssessment captures "
        "assessment submissions and scores. studentVle records detailed Virtual Learning Environment (VLE) activity at the student level. "
        "These raw tables will not be modeled directly. Instead, Milestone 4 will require aggregation of assessment and Virtual Learning Environment behaviour into student-level features. "
        "The join logic is centered on code_module, code_presentation, and id_student."
    )
    add_table(
        doc,
        ["OULAD Table", "Rows", "Planned Use in Project"],
        [
            ["studentInfo", "32,593", "Core student attributes and final_result target"],
            ["studentRegistration", "32,593", "Registration and withdrawal timing features"],
            ["studentAssessment", "173,912", "Assessment participation and performance aggregates"],
            ["studentVle", "10,655,280", "Student Virtual Learning Environment engagement features after aggregation"],
            ["assessments", "206", "Assessment metadata for feature enrichment"],
            ["courses", "22", "Course/module reference"],
            ["vle", "6,364", "Virtual Learning Environment activity metadata if needed for interpretation"],
        ],
    )
    add_para(
        doc,
        "The main dataset limitation at this stage is not access but analytical structure. UCI is easier to interpret and "
        "model quickly, but some of its strongest predictors may arrive too late for a genuinely early intervention system. "
        "OULAD offers richer early behavioural signals, but it requires more extensive joins, aggregation, and leakage control."
    )

    add_heading(doc, "3. Exploratory Data Analysis (EDA)", level=1)
    add_heading(doc, "3.1 Statistical Summary and Initial EDA - UCI Dataset", level=2)
    add_para(
        doc,
        f"The UCI dataset contains {uci_total} records and 37 columns, of which 36 are predictor variables and one is the "
        "multiclass target variable. The first checks confirmed zero missing values and zero duplicate rows. The class "
        f"distribution is moderately imbalanced: Graduate = {uci_target['Graduate']} ({pct(uci_target['Graduate'], uci_total)}%), "
        f"Dropout = {uci_target['Dropout']} ({pct(uci_target['Dropout'], uci_total)}%), and Enrolled = {uci_target['Enrolled']} "
        f"({pct(uci_target['Enrolled'], uci_total)}%). This means the Enrolled class is the smallest group and may require special "
        "attention during later modeling and evaluation."
    )
    add_para(
        doc,
        "Next, group-level descriptive statistics were compared across the three outcomes. Students in the Graduate group averaged "
        "6.23 approved first-semester units and 6.18 approved second-semester units, compared with 2.55 and 1.94 respectively "
        "in the Dropout group. Graduate students also had higher semester grades and higher scholarship-holder rates, while dropout "
        "students had higher debtor rates and lower tuition-payment compliance. These differences suggest that academic progress and "
        "financial standing are strong predictive signals in the benchmark data."
    )
    add_para(
        doc,
        "Correlation analysis supported the same pattern. The strongest correlations with a binary dropout indicator were "
        "Curricular units 2nd sem (grade) (-0.572), Curricular units 2nd sem (approved) (-0.570), Curricular units 1st sem "
        "(grade) (-0.481), Curricular units 1st sem (approved) (-0.479), and Age at enrollment (0.254). These findings suggest "
        "that academic throughput and age-related enrollment patterns may matter more than macroeconomic background fields such as "
        "GDP, inflation, or unemployment in this dataset."
    )
    add_para(
        doc,
        "Outlier review using the IQR rule identified notable counts in age and grade-related variables, including 441 observations "
        "in Age at enrollment and more than 700 observations in each semester-grade field. At this stage, these values are treated "
        "as potentially meaningful heterogeneity rather than automatic data errors, because they likely capture real differences in "
        "student pathways, course load, and academic performance."
    )

    add_heading(doc, "3.2 Statistical Summary and Initial EDA - OULAD", level=2)
    add_para(
        doc,
        f"OULAD is structurally richer than UCI, so it was reviewed at both table level and merged-feature level. The core "
        f"studentInfo table contains {oulad_total} student records and 12 variables. The target distribution is Pass = {oulad_target['Pass']} "
        f"({pct(oulad_target['Pass'], oulad_total)}%), Withdrawn = {oulad_target['Withdrawn']} ({pct(oulad_target['Withdrawn'], oulad_total)}%), "
        f"Fail = {oulad_target['Fail']} ({pct(oulad_target['Fail'], oulad_total)}%), and Distinction = {oulad_target['Distinction']} "
        f"({pct(oulad_target['Distinction'], oulad_total)}%). This produces a sizable withdrawal class and confirms that OULAD is suitable "
        "for attrition-oriented analysis."
    )
    add_para(
        doc,
        "Initial data quality inspection found no duplicate rows in studentInfo, studentRegistration, or studentAssessment. However, "
        "important missing-value patterns are present. In studentInfo, imd_band has 1,111 missing values (3.41%). In studentRegistration, "
        "date_registration has 45 missing values (0.14%), while date_unregistration has 22,521 missing values (69.10%), which is largely "
        "structural because most students did not withdraw. In studentAssessment, score has 173 missing values (0.10%). These findings are "
        "not severe enough to make the dataset unusable, but they will directly influence preprocessing decisions."
    )
    add_para(
        doc,
        "After the key OULAD tables were merged at student level, several useful early warning signals became clear. Students with Distinction averaged 8.71 assessment "
        "records, a mean score of 88.49, and 2,666.76 total Virtual Learning Environment clicks. The Withdrawn group averaged only 1.29 assessment records, a mean "
        "score of 29.50, and 313.95 total Virtual Learning Environment clicks. The Pass group remained much closer to Distinction than to Withdrawn on both engagement "
        "and assessment measures. One interesting finding is that Withdrawn students averaged 91.43 studied credits, which is higher than all "
        "other groups. This suggests that withdrawal risk may not only be associated with low engagement, but may also reflect course-load intensity "
        "or challenging study combinations."
    )
    add_para(
        doc,
        "The studentVle table required aggregation to become analytically useful. After grouping event-level Virtual Learning Environment records into student-level behaviour "
        "features, the resulting aggregated table contained 29,228 grouped student records. This confirmed that OULAD is computationally tractable "
        "on local hardware when event logs are processed in chunks and reduced to modelling-friendly summaries."
    )

    add_heading(doc, "3.3 Data Visualization and Interpretation", level=2)
    add_figure(doc, uci["figures"]["target_distribution"], "Figure 1. UCI target distribution shows moderate class imbalance, with Enrolled as the smallest class.")
    add_para(
        doc,
        "Figure 1 confirms that the UCI target variable is not severely imbalanced, but the smaller Enrolled class may still make multiclass "
        "classification more difficult than a binary attrition task."
    )
    add_figure(doc, uci["figures"]["approved_units_boxplot"], "Figure 2. UCI first-semester approved curricular units differ sharply across outcomes.")
    add_para(
        doc,
        "Figure 2 shows the separation between dropout and non-dropout outcomes more clearly than a simple mean comparison. Lower approved-unit "
        "counts appear strongly associated with the Dropout class, which supports the use of approved curricular units as an important benchmark feature."
    )
    add_figure(doc, uci["figures"]["correlation_heatmap"], "Figure 3. UCI selected-variable heatmap shows strong positive relationships among semester performance variables.")
    add_para(
        doc,
        "The heatmap shows that several academic progress variables move closely together, especially across first- and second-semester approved units and grades. "
        "This is useful for prediction, but it also means baseline models will need some attention to feature selection and multicollinearity."
    )
    add_figure(doc, oulad["figures"]["target_distribution"], "Figure 4. OULAD final_result distribution supports a strong attrition-oriented evaluation setting.")
    add_para(
        doc,
        "Figure 4 highlights the practical value of OULAD for this project: the Withdrawn class is large enough to support meaningful early warning analysis, "
        "while the Pass, Fail, and Distinction groups provide richer outcome variation than a purely binary benchmark."
    )
    add_figure(doc, oulad["figures"]["studied_credits_boxplot"], "Figure 5. OULAD studied credits vary across final outcome groups, with the Withdrawn group carrying the highest average load.")
    add_para(
        doc,
        "This visualization suggests that studied_credits may be a useful contextual feature rather than a direct risk proxy. Higher credit load does not imply "
        "success in this dataset and may instead capture higher demand or more fragile student workloads."
    )
    add_figure(doc, oulad["figures"]["log_clicks_boxplot"], "Figure 6. OULAD log total Virtual Learning Environment clicks show that withdrawn students are substantially less active online.")
    add_para(
        doc,
        "Figure 6 is one of the clearest findings in the milestone. The Withdrawn group has much lower Virtual Learning Environment activity than Pass and Distinction, "
        "which supports the use of early behavioural engagement as a central feature family for later modeling."
    )
    add_figure(doc, oulad["figures"]["mean_score_boxplot"], "Figure 7. OULAD mean assessment score also separates Withdrawn students from the more successful groups.")
    add_para(
        doc,
        "Figure 7 suggests that academic performance and behavioural engagement reinforce one another in OULAD. Students who later withdraw do not simply log in less; "
        "they also tend to submit fewer assessments and perform worse on the assessments they do complete."
    )

    add_heading(doc, "3.4 Initial Business Insights and Analytical Interpretation", level=2)
    add_para(
        doc,
        "Several early business insights emerged from the analysis. First, both datasets support the core business problem: identifying at-risk students before "
        "attrition becomes final. In UCI, financial and academic progress fields appear especially relevant to student persistence. In OULAD, the clearest signals are "
        "assessment participation, assessment performance, and behavioural engagement. This means that institutions using different data environments may need different "
        "feature strategies even when the business objective is similar."
    )
    add_para(
        doc,
        "Second, the analysis suggests an important practical distinction between benchmark accuracy and intervention timing. The UCI dataset includes very strong "
        "second-semester performance fields, but those variables may be too late to support genuinely early action. By contrast, OULAD provides behaviour streams that "
        "can potentially be truncated into earlier time windows, making it a more realistic environment for early warning design."
    )
    add_para(
        doc,
        "Third, the observed gap between Withdrawn students and the other OULAD groups suggests that a harmonized attrition outcome for later cross-dataset testing is "
        "defensible. The behavioural and assessment differences are large enough to justify later experiments that compare persistence versus attrition across the two datasets."
    )

    add_heading(doc, "4. Data Quality and Preprocessing Assessment", level=1)
    add_heading(doc, "4.1 Data Quality Concerns", level=2)
    add_para(
        doc,
        "Overall, both datasets are usable, but they differ strongly in preparation burden. UCI is immediately usable for baseline modeling because it has no missing "
        "values, no duplicates, and a compact single-table structure. The main quality concern is semantic rather than structural: several numeric-coded fields should "
        "be treated as categorical or ordinal rather than continuous. In addition, some strong academic variables may create timing concerns if the objective is truly "
        "early intervention rather than benchmark replication."
    )
    add_para(
        doc,
        "OULAD presents more realistic data quality challenges. The missingness in imd_band and date_registration is limited and manageable, but the missingness in "
        "date_unregistration is primarily structural and cannot be treated as ordinary random missing data. The studentAssessment score field also contains a small number "
        "of missing values, and the very large studentVle table introduces scale, aggregation, and consistency challenges. These are not reasons to reject the dataset; "
        "they are precisely the reasons it is valuable as the primary evaluation dataset."
    )

    add_heading(doc, "4.2 Preliminary Preprocessing Plan", level=2)
    add_bullet(doc, "For UCI, preserve the full feature set for reproduction of the base paper, but separately flag later-semester variables when discussing early warning realism.")
    add_bullet(doc, "Convert code-based fields in UCI to appropriate categorical or ordinal treatment where needed; use scaling selectively for linear baselines such as Logistic Regression.")
    add_bullet(doc, "For OULAD, merge studentInfo, studentRegistration, studentAssessment, and aggregated studentVle records at the student-module-presentation level.")
    add_bullet(doc, "Impute or otherwise handle missing imd_band values, review the small number of missing scores, and exclude date_unregistration from predictive inputs to prevent leakage.")
    add_bullet(doc, "Create student-level behavioural features from studentVle using chunked aggregation of Virtual Learning Environment activity, then define time-aware cutoffs for early-warning feature windows.")
    add_bullet(doc, "For later cross-dataset experiments, harmonize labels into attrition versus persistence and align features by shared families rather than by raw column names.")

    add_heading(doc, "4.3 Anticipated Transformation and Feature Engineering Challenges", level=2)
    add_para(
        doc,
        "The most important preprocessing challenge is leakage control. In UCI, later academic variables may artificially strengthen predictions while weakening real-world intervention value. "
        "In OULAD, date_unregistration and late-course behaviours could reveal the outcome too directly if included without temporal cutoffs. The second major challenge is feature harmonization. "
        "The two datasets do not share the same schema, so cross-dataset modeling will require alignment by concept families such as demographics, prior preparation, academic progression, and early engagement."
    )
    add_para(
        doc,
        "A third challenge is scale. Although OULAD is computationally feasible on local hardware when processed in chunks, its event-level data are not modeling-ready in raw form. "
        "Aggregation and reduction are therefore essential parts of preprocessing rather than optional enhancements."
    )

    add_heading(doc, "5. Modeling Readiness Assessment", level=1)
    add_para(
        doc,
        "The dataset readiness assessment is positive overall. UCI is immediately ready for baseline reproduction modeling in Milestone 4, subject only to normal preprocessing decisions such "
        "as train-test splitting, feature handling, and multiclass evaluation design. OULAD is also viable for modeling, but it is not yet row-level ready in the same way because studentAssessment "
        "and studentVle must first be aggregated into student-level Virtual Learning Environment features and filtered to avoid leakage."
    )
    add_para(
        doc,
        "From a computational perspective, the project remains feasible within standard laptop or Google Colab constraints. The largest raw table is studentVle, but this milestone confirmed that "
        "chunk-based aggregation is sufficient for reducing it into useful behavioural summaries. Overfitting remains a risk in both datasets, though for different reasons: UCI contains strong "
        "later-semester variables, while OULAD may generate many engineered features from behaviour logs. Both issues can be managed through careful feature design, baseline comparison, and explicit validation."
    )
    add_para(
        doc,
        "The main risks entering Milestone 4 are therefore methodological rather than access-related. These include deciding on fair early-warning time windows, controlling label leakage, handling "
        "class imbalance, and building a cross-dataset representation that is realistic enough for later portability testing. Even with those risks, the current milestone confirms that the project is "
        "ready to move into preprocessing and baseline modeling."
    )

    add_heading(doc, "6. References", level=1)
    refs = [
        '[1] M. M. Islam, F. H. Sojib, M. F. H. Mihad, M. Hasan, and M. Rahman, "The integration of explainable AI in Educational Data Mining for student academic performance prediction and support system," Telematics and Informatics Reports, vol. 18, Art. no. 100203, 2025, doi: 10.1016/j.teler.2025.100203.',
        '[2] V. Realinho, J. Machado, L. Baptista, and M. V. Martins, "Predicting Student Dropout and Academic Success," Data, vol. 7, no. 11, Art. no. 146, 2022, doi: 10.3390/data7110146.',
        '[3] V. Realinho, M. V. Martins, J. Machado, and L. Baptista, Predict Students\' Dropout and Academic Success [Dataset], UCI Machine Learning Repository, 2021. doi: 10.24432/C5MC89.',
        '[4] J. Kuzilek, M. Hlosta, and Z. Zdrahal, "Open University Learning Analytics dataset," Scientific Data, vol. 4, Art. no. 170171, 2017, doi: 10.1038/sdata.2017.171.',
        '[5] C. R. Cirak, H. Akilli, and Y. Ekinci, "Development of an early warning system for higher education institutions by predicting first-year student academic performance," Higher Education Quarterly, vol. 78, no. 4, 2024, doi: 10.1111/hequ.12539.',
        '[6] M. Phan, A. De Caigny, and K. Coussement, "A decision support framework to incorporate textual data for early student dropout prediction in higher education," Decision Support Systems, vol. 168, Art. no. 113940, 2023, doi: 10.1016/j.dss.2023.113940.',
    ]
    for ref in refs:
        add_para(doc, ref, after=3, line=1.0)

    doc.save(WORK_DOC)
    doc.save(ALT_WORK_DOC)


if __name__ == "__main__":
    build()
